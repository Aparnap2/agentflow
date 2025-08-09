"""
Document Converter Service - Converts between different document formats
"""
import os
import io
import tempfile
from typing import Dict, Any, Optional, Literal
from pathlib import Path
from datetime import datetime
import markdown2
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt
from loguru import logger

class DocumentConverter:
    """Converts documents between different formats"""
    
    def __init__(self):
        """Initialize document converter"""
        self.output_dir = Path("./data/documents")
        self.output_dir.mkdir(exist_ok=True, parents=True)
        
        # Create subdirectories for different document types
        for doc_type in ["marketing", "finance", "sales", "legal", "operations", "research"]:
            (self.output_dir / doc_type).mkdir(exist_ok=True)
    
    def markdown_to_pdf(self, markdown_content: str, metadata: Dict[str, Any] = None) -> bytes:
        """Convert Markdown to PDF using ReportLab"""
        try:
            # Convert markdown to HTML
            html_content = markdown2.markdown(
                markdown_content,
                extras=["tables", "code-friendly", "cuddled-lists", "fenced-code-blocks"]
            )
            
            # Create PDF buffer
            buffer = io.BytesIO()
            
            # Create PDF document
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            
            # Add custom styles
            styles.add(ParagraphStyle(
                name='Heading1',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=12
            ))
            styles.add(ParagraphStyle(
                name='Heading2',
                parent=styles['Heading2'],
                fontSize=16,
                spaceAfter=10
            ))
            styles.add(ParagraphStyle(
                name='Heading3',
                parent=styles['Heading3'],
                fontSize=14,
                spaceAfter=8
            ))
            
            # Parse HTML and create PDF elements
            elements = []
            
            # Add metadata if provided
            if metadata:
                if "title" in metadata:
                    elements.append(Paragraph(metadata["title"], styles["Title"]))
                    elements.append(Spacer(1, 12))
                
                if "author" in metadata:
                    elements.append(Paragraph(f"Author: {metadata['author']}", styles["Normal"]))
                    elements.append(Spacer(1, 12))
                
                if "date" in metadata:
                    elements.append(Paragraph(f"Date: {metadata['date']}", styles["Normal"]))
                    elements.append(Spacer(1, 12))
            
            # Simple HTML parsing (this is a basic implementation)
            # For a more robust solution, consider using an HTML parser
            lines = html_content.split("\n")
            for line in lines:
                if line.strip():
                    if line.startswith("<h1>"):
                        text = line.replace("<h1>", "").replace("</h1>", "")
                        elements.append(Paragraph(text, styles["Heading1"]))
                    elif line.startswith("<h2>"):
                        text = line.replace("<h2>", "").replace("</h2>", "")
                        elements.append(Paragraph(text, styles["Heading2"]))
                    elif line.startswith("<h3>"):
                        text = line.replace("<h3>", "").replace("</h3>", "")
                        elements.append(Paragraph(text, styles["Heading3"]))
                    elif line.startswith("<p>"):
                        text = line.replace("<p>", "").replace("</p>", "")
                        elements.append(Paragraph(text, styles["Normal"]))
                    else:
                        # Default to normal text
                        elements.append(Paragraph(line, styles["Normal"]))
                    
                    elements.append(Spacer(1, 6))
            
            # Build PDF
            doc.build(elements)
            
            # Get PDF content
            pdf_content = buffer.getvalue()
            buffer.close()
            
            return pdf_content
            
        except Exception as e:
            logger.error(f"Failed to convert Markdown to PDF: {e}")
            raise ValueError(f"Failed to convert Markdown to PDF: {e}")
    
    def markdown_to_docx(self, markdown_content: str, metadata: Dict[str, Any] = None) -> bytes:
        """Convert Markdown to DOCX using python-docx"""
        try:
            # Create Word document
            doc = Document()
            
            # Add metadata if provided
            if metadata:
                if "title" in metadata:
                    doc.add_heading(metadata["title"], 0)
                
                if "author" in metadata or "date" in metadata:
                    meta_text = ""
                    if "author" in metadata:
                        meta_text += f"Author: {metadata['author']}"
                    if "date" in metadata:
                        if meta_text:
                            meta_text += " | "
                        meta_text += f"Date: {metadata['date']}"
                    
                    doc.add_paragraph(meta_text)
                    doc.add_paragraph()  # Empty paragraph for spacing
            
            # Parse markdown (simple implementation)
            lines = markdown_content.split("\n")
            current_list = []
            in_list = False
            
            for line in lines:
                # Handle headings
                if line.startswith("# "):
                    doc.add_heading(line[2:], 1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], 2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], 3)
                # Handle lists
                elif line.startswith("- ") or line.startswith("* "):
                    if not in_list:
                        in_list = True
                        current_list = []
                    
                    current_list.append(line[2:])
                # Handle normal text
                else:
                    # First flush any pending list
                    if in_list:
                        for item in current_list:
                            doc.add_paragraph(item, style='ListBullet')
                        current_list = []
                        in_list = False
                    
                    # Add normal paragraph if line is not empty
                    if line.strip():
                        doc.add_paragraph(line)
            
            # Flush any remaining list items
            if in_list:
                for item in current_list:
                    doc.add_paragraph(item, style='ListBullet')
            
            # Save to bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            docx_content = buffer.getvalue()
            buffer.close()
            
            return docx_content
            
        except Exception as e:
            logger.error(f"Failed to convert Markdown to DOCX: {e}")
            raise ValueError(f"Failed to convert Markdown to DOCX: {e}")
    
    def markdown_to_html(self, markdown_content: str, metadata: Dict[str, Any] = None) -> str:
        """Convert Markdown to HTML"""
        try:
            # Convert markdown to HTML
            html_content = markdown2.markdown(
                markdown_content,
                extras=["tables", "code-friendly", "cuddled-lists", "fenced-code-blocks"]
            )
            
            # Add metadata if provided
            if metadata:
                header = "<header>"
                if "title" in metadata:
                    header += f"<h1>{metadata['title']}</h1>"
                
                if "author" in metadata or "date" in metadata:
                    header += "<p>"
                    if "author" in metadata:
                        header += f"Author: {metadata['author']}"
                    if "date" in metadata:
                        if "author" in metadata:
                            header += " | "
                        header += f"Date: {metadata['date']}"
                    header += "</p>"
                
                header += "</header>"
                html_content = f"{header}\n{html_content}"
            
            # Wrap in basic HTML structure
            full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{metadata.get('title', 'Document') if metadata else 'Document'}</title>
    <style>
        body {{ font-family: Arial, sans-serif; line-height: 1.6; max-width: 800px; margin: 0 auto; padding: 20px; }}
        h1, h2, h3 {{ color: #333; }}
        table {{ border-collapse: collapse; width: 100%; }}
        table, th, td {{ border: 1px solid #ddd; padding: 8px; }}
        th {{ background-color: #f2f2f2; }}
        code {{ background-color: #f5f5f5; padding: 2px 4px; border-radius: 4px; }}
        pre {{ background-color: #f5f5f5; padding: 10px; border-radius: 4px; overflow-x: auto; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""
            
            return full_html
            
        except Exception as e:
            logger.error(f"Failed to convert Markdown to HTML: {e}")
            raise ValueError(f"Failed to convert Markdown to HTML: {e}")
    
    def save_document(self, content: bytes, doc_type: str, filename: str, 
                     format: Literal["pdf", "docx", "md", "html", "txt"]) -> str:
        """Save document to file system"""
        try:
            # Ensure filename has correct extension
            if not filename.endswith(f".{format}"):
                filename = f"{filename}.{format}"
            
            # Create full path
            file_path = self.output_dir / doc_type / filename
            
            # Write content
            with open(file_path, 'wb') as f:
                f.write(content)
            
            logger.info(f"Saved document: {file_path}")
            return str(file_path)
            
        except Exception as e:
            logger.error(f"Failed to save document: {e}")
            raise ValueError(f"Failed to save document: {e}")
    
    def convert_document(self, content: str, source_format: str, target_format: str, 
                        metadata: Dict[str, Any] = None) -> bytes:
        """Convert document from one format to another"""
        try:
            # Handle different conversion paths
            if source_format == "md" and target_format == "pdf":
                return self.markdown_to_pdf(content, metadata)
            elif source_format == "md" and target_format == "docx":
                return self.markdown_to_docx(content, metadata)
            elif source_format == "md" and target_format == "html":
                return self.markdown_to_html(content, metadata).encode('utf-8')
            elif source_format == target_format:
                # No conversion needed
                return content.encode('utf-8')
            else:
                raise ValueError(f"Unsupported conversion: {source_format} to {target_format}")
                
        except Exception as e:
            logger.error(f"Failed to convert document: {e}")
            raise ValueError(f"Failed to convert document: {e}")
    
    def generate_document(self, content: str, doc_type: str, filename: str,
                         format: Literal["pdf", "docx", "md", "html", "txt"],
                         metadata: Dict[str, Any] = None) -> Dict[str, Any]:
        """Generate document in specified format and save it"""
        try:
            # Determine source format (assume markdown if not specified)
            source_format = "md"
            
            # Convert document
            doc_content = self.convert_document(content, source_format, format, metadata)
            
            # Save document
            file_path = self.save_document(doc_content, doc_type, filename, format)
            
            # Return document info
            return {
                "path": file_path,
                "type": doc_type,
                "format": format,
                "size": len(doc_content),
                "created_at": datetime.now().isoformat(),
                "metadata": metadata or {}
            }
            
        except Exception as e:
            logger.error(f"Failed to generate document: {e}")
            raise ValueError(f"Failed to generate document: {e}")

# Global instance
document_converter = DocumentConverter()